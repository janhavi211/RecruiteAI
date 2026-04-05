#!/usr/bin/env python3
"""
RecruitAI - Advanced Resume Intelligence Engine v3
Fixed: Multiple PDF extraction methods, better skill matching,
       robust fallbacks, detailed debug logging
"""

import sys
import json
import re
import os
from datetime import datetime

# ══════════════════════════════════════════════════════════════
#  DEBUG HELPER - writes to stderr so it doesn't break JSON output
# ══════════════════════════════════════════════════════════════
def debug(msg):
    pass
# ══════════════════════════════════════════════════════════════
#  SKILLS DATABASE (flat list for easier matching)
# ══════════════════════════════════════════════════════════════
SKILLS_DB = {
    # Programming
    "python":{"cat":"Programming","w":9},"java":{"cat":"Programming","w":9},
    "javascript":{"cat":"Programming","w":9},"js":{"cat":"Programming","w":8},
    "typescript":{"cat":"Programming","w":8},"ts":{"cat":"Programming","w":7},
    "c++":{"cat":"Programming","w":9},"cpp":{"cat":"Programming","w":9},
    "c#":{"cat":"Programming","w":8},"csharp":{"cat":"Programming","w":8},
    "php":{"cat":"Programming","w":8},"ruby":{"cat":"Programming","w":7},
    "go":{"cat":"Programming","w":8},"golang":{"cat":"Programming","w":8},
    "rust":{"cat":"Programming","w":8},"kotlin":{"cat":"Programming","w":8},
    "swift":{"cat":"Programming","w":8},"scala":{"cat":"Programming","w":7},
    "r":{"cat":"Programming","w":7},"matlab":{"cat":"Programming","w":6},
    "perl":{"cat":"Programming","w":5},"dart":{"cat":"Programming","w":7},
    "bash":{"cat":"Scripting","w":7},"shell":{"cat":"Scripting","w":7},
    "powershell":{"cat":"Scripting","w":6},
    # Frontend
    "html":{"cat":"Frontend","w":7},"css":{"cat":"Frontend","w":7},
    "html5":{"cat":"Frontend","w":7},"css3":{"cat":"Frontend","w":7},
    "react":{"cat":"Frontend","w":9},"reactjs":{"cat":"Frontend","w":9},
    "vue":{"cat":"Frontend","w":8},"vuejs":{"cat":"Frontend","w":8},
    "angular":{"cat":"Frontend","w":8},"angularjs":{"cat":"Frontend","w":7},
    "svelte":{"cat":"Frontend","w":7},"jquery":{"cat":"Frontend","w":6},
    "bootstrap":{"cat":"Frontend","w":6},"tailwind":{"cat":"Frontend","w":7},
    "tailwindcss":{"cat":"Frontend","w":7},"sass":{"cat":"Frontend","w":6},
    "scss":{"cat":"Frontend","w":6},"webpack":{"cat":"Frontend","w":7},
    "vite":{"cat":"Frontend","w":7},"redux":{"cat":"Frontend","w":7},
    "graphql":{"cat":"Frontend","w":8},"nextjs":{"cat":"Frontend","w":9},
    "gatsby":{"cat":"Frontend","w":7},"nuxt":{"cat":"Frontend","w":7},
    # Backend
    "nodejs":{"cat":"Backend","w":9},"express":{"cat":"Backend","w":8},
    "expressjs":{"cat":"Backend","w":8},"django":{"cat":"Backend","w":9},
    "flask":{"cat":"Backend","w":8},"fastapi":{"cat":"Backend","w":8},
    "spring":{"cat":"Backend","w":8},"springboot":{"cat":"Backend","w":9},
    "laravel":{"cat":"Backend","w":8},"codeigniter":{"cat":"Backend","w":6},
    "rails":{"cat":"Backend","w":7},"nestjs":{"cat":"Backend","w":8},
    "hibernate":{"cat":"Backend","w":6},"strapi":{"cat":"Backend","w":6},
    # Database
    "mysql":{"cat":"Database","w":8},"postgresql":{"cat":"Database","w":9},
    "postgres":{"cat":"Database","w":9},"mongodb":{"cat":"Database","w":8},
    "mongo":{"cat":"Database","w":8},"redis":{"cat":"Database","w":8},
    "sqlite":{"cat":"Database","w":6},"oracle":{"cat":"Database","w":7},
    "mssql":{"cat":"Database","w":7},"sqlserver":{"cat":"Database","w":7},
    "cassandra":{"cat":"Database","w":7},"elasticsearch":{"cat":"Database","w":8},
    "firebase":{"cat":"Database","w":7},"dynamodb":{"cat":"Database","w":7},
    "mariadb":{"cat":"Database","w":6},"neo4j":{"cat":"Database","w":7},
    "supabase":{"cat":"Database","w":7},"sql":{"cat":"Database","w":8},
    "nosql":{"cat":"Database","w":7},"prisma":{"cat":"Database","w":7},
    "sequelize":{"cat":"Database","w":6},"mongoose":{"cat":"Database","w":6},
    # Cloud / DevOps
    "aws":{"cat":"Cloud","w":10},"azure":{"cat":"Cloud","w":10},
    "gcp":{"cat":"Cloud","w":10},"heroku":{"cat":"Cloud","w":6},
    "vercel":{"cat":"Cloud","w":7},"netlify":{"cat":"Cloud","w":6},
    "docker":{"cat":"DevOps","w":9},"kubernetes":{"cat":"DevOps","w":9},
    "k8s":{"cat":"DevOps","w":9},"jenkins":{"cat":"DevOps","w":8},
    "terraform":{"cat":"DevOps","w":8},"ansible":{"cat":"DevOps","w":7},
    "nginx":{"cat":"DevOps","w":7},"apache":{"cat":"DevOps","w":6},
    "linux":{"cat":"System","w":8},"ubuntu":{"cat":"System","w":7},
    "unix":{"cat":"System","w":7},"windows":{"cat":"System","w":5},
    "git":{"cat":"VCS","w":9},"github":{"cat":"VCS","w":8},
    "gitlab":{"cat":"VCS","w":7},"bitbucket":{"cat":"VCS","w":6},
    "svn":{"cat":"VCS","w":4},
    # AI / ML / Data
    "tensorflow":{"cat":"AI/ML","w":9},"pytorch":{"cat":"AI/ML","w":9},
    "keras":{"cat":"AI/ML","w":8},"opencv":{"cat":"AI/ML","w":8},
    "sklearn":{"cat":"AI/ML","w":8},"scipy":{"cat":"AI/ML","w":7},
    "pandas":{"cat":"Data Science","w":8},"numpy":{"cat":"Data Science","w":8},
    "matplotlib":{"cat":"Data Science","w":7},"seaborn":{"cat":"Data Science","w":6},
    "jupyter":{"cat":"Data Science","w":7},"spark":{"cat":"Big Data","w":9},
    "hadoop":{"cat":"Big Data","w":7},"kafka":{"cat":"Messaging","w":8},
    "tableau":{"cat":"Data Viz","w":8},"powerbi":{"cat":"Data Viz","w":8},
    "langchain":{"cat":"AI/ML","w":9},"openai":{"cat":"AI/ML","w":8},
    "huggingface":{"cat":"AI/ML","w":9},"llm":{"cat":"AI/ML","w":9},
    # Mobile
    "android":{"cat":"Mobile","w":8},"ios":{"cat":"Mobile","w":8},
    "flutter":{"cat":"Mobile","w":9},"kotlin":{"cat":"Mobile","w":8},
    "swift":{"cat":"Mobile","w":8},"ionic":{"cat":"Mobile","w":6},
    "xamarin":{"cat":"Mobile","w":6},
    # Design
    "figma":{"cat":"Design","w":8},"adobexd":{"cat":"Design","w":7},
    "sketch":{"cat":"Design","w":6},"photoshop":{"cat":"Design","w":7},
    "illustrator":{"cat":"Design","w":7},"canva":{"cat":"Design","w":5},
    # Testing
    "selenium":{"cat":"Testing","w":7},"jest":{"cat":"Testing","w":7},
    "pytest":{"cat":"Testing","w":7},"junit":{"cat":"Testing","w":7},
    "cypress":{"cat":"Testing","w":8},"playwright":{"cat":"Testing","w":8},
    "postman":{"cat":"Testing","w":7},"mocha":{"cat":"Testing","w":6},
    "appium":{"cat":"Testing","w":6},"testng":{"cat":"Testing","w":6},
    # Management
    "agile":{"cat":"Management","w":8},"scrum":{"cat":"Management","w":8},
    "kanban":{"cat":"Management","w":7},"jira":{"cat":"Management","w":7},
    "confluence":{"cat":"Management","w":6},"trello":{"cat":"Management","w":5},
    "pmp":{"cat":"Management","w":9},"prince2":{"cat":"Management","w":8},
    # Architecture
    "restapi":{"cat":"Architecture","w":8},"rest":{"cat":"Architecture","w":7},
    "microservices":{"cat":"Architecture","w":9},"mvc":{"cat":"Architecture","w":7},
    "oop":{"cat":"Architecture","w":7},"api":{"cat":"Architecture","w":7},
    "serverless":{"cat":"Architecture","w":8},"grpc":{"cat":"Architecture","w":7},
    "soap":{"cat":"Architecture","w":5},"graphql":{"cat":"Architecture","w":8},
    # Security
    "cybersecurity":{"cat":"Security","w":9},"owasp":{"cat":"Security","w":8},
    "jwt":{"cat":"Security","w":7},"oauth":{"cat":"Security","w":7},
    "ssl":{"cat":"Security","w":6},"encryption":{"cat":"Security","w":7},
    # Office
    "excel":{"cat":"Productivity","w":6},"word":{"cat":"Productivity","w":5},
    "powerpoint":{"cat":"Productivity","w":5},"msoffice":{"cat":"Productivity","w":5},
    # Blockchain
    "blockchain":{"cat":"Blockchain","w":8},"solidity":{"cat":"Blockchain","w":8},
    "ethereum":{"cat":"Blockchain","w":7},"web3":{"cat":"Blockchain","w":8},
    # Soft skills
    "leadership":{"cat":"Soft Skills","w":8},"communication":{"cat":"Soft Skills","w":7},
    "teamwork":{"cat":"Soft Skills","w":7},"management":{"cat":"Soft Skills","w":7},
}

# Multi-word skills to check as phrases
MULTI_WORD_SKILLS = {
    "machine learning":{"cat":"AI/ML","w":10,"display":"Machine Learning"},
    "deep learning":{"cat":"AI/ML","w":10,"display":"Deep Learning"},
    "artificial intelligence":{"cat":"AI/ML","w":10,"display":"Artificial Intelligence"},
    "natural language processing":{"cat":"AI/ML","w":9,"display":"NLP"},
    "computer vision":{"cat":"AI/ML","w":9,"display":"Computer Vision"},
    "neural network":{"cat":"AI/ML","w":9,"display":"Neural Networks"},
    "neural networks":{"cat":"AI/ML","w":9,"display":"Neural Networks"},
    "data science":{"cat":"Data Science","w":9,"display":"Data Science"},
    "data analysis":{"cat":"Data Science","w":8,"display":"Data Analysis"},
    "data analytics":{"cat":"Data Science","w":8,"display":"Data Analytics"},
    "data structures":{"cat":"Programming","w":8,"display":"Data Structures"},
    "operating system":{"cat":"System","w":6,"display":"Operating Systems"},
    "operating systems":{"cat":"System","w":6,"display":"Operating Systems"},
    "object oriented":{"cat":"Programming","w":7,"display":"OOP"},
    "object-oriented":{"cat":"Programming","w":7,"display":"OOP"},
    "node.js":{"cat":"Backend","w":9,"display":"Node.js"},
    "next.js":{"cat":"Frontend","w":9,"display":"Next.js"},
    "react.js":{"cat":"Frontend","w":9,"display":"React.js"},
    "vue.js":{"cat":"Frontend","w":8,"display":"Vue.js"},
    "express.js":{"cat":"Backend","w":8,"display":"Express.js"},
    "spring boot":{"cat":"Backend","w":9,"display":"Spring Boot"},
    "asp.net":{"cat":"Backend","w":8,"display":"ASP.NET"},
    "react native":{"cat":"Mobile","w":9,"display":"React Native"},
    "power bi":{"cat":"Data Viz","w":8,"display":"Power BI"},
    "ms office":{"cat":"Productivity","w":5,"display":"MS Office"},
    "microsoft office":{"cat":"Productivity","w":5,"display":"MS Office"},
    "adobe xd":{"cat":"Design","w":7,"display":"Adobe XD"},
    "scikit learn":{"cat":"AI/ML","w":8,"display":"Scikit-Learn"},
    "scikit-learn":{"cat":"AI/ML","w":8,"display":"Scikit-Learn"},
    "google cloud":{"cat":"Cloud","w":10,"display":"Google Cloud"},
    "github actions":{"cat":"DevOps","w":8,"display":"GitHub Actions"},
    "rest api":{"cat":"Architecture","w":8,"display":"REST API"},
    "restful api":{"cat":"Architecture","w":8,"display":"RESTful API"},
    "user interface":{"cat":"Design","w":6,"display":"UI Design"},
    "user experience":{"cat":"Design","w":7,"display":"UX Design"},
    "problem solving":{"cat":"Soft Skills","w":8,"display":"Problem Solving"},
    "project management":{"cat":"Management","w":8,"display":"Project Management"},
    "version control":{"cat":"VCS","w":7,"display":"Version Control"},
    "ci/cd":{"cat":"DevOps","w":8,"display":"CI/CD"},
    "ci cd":{"cat":"DevOps","w":8,"display":"CI/CD"},
    "smart contracts":{"cat":"Blockchain","w":8,"display":"Smart Contracts"},
    "hugging face":{"cat":"AI/ML","w":9,"display":"Hugging Face"},
    "sql server":{"cat":"Database","w":7,"display":"SQL Server"},
    "jetpack compose":{"cat":"Mobile","w":8,"display":"Jetpack Compose"},
    "swift ui":{"cat":"Mobile","w":7,"display":"SwiftUI"},
    "unit testing":{"cat":"Testing","w":7,"display":"Unit Testing"},
    "integration testing":{"cat":"Testing","w":7,"display":"Integration Testing"},
    "ethical hacking":{"cat":"Security","w":9,"display":"Ethical Hacking"},
    "penetration testing":{"cat":"Security","w":9,"display":"Penetration Testing"},
    "d3.js":{"cat":"Frontend","w":7,"display":"D3.js"},
    "three.js":{"cat":"Frontend","w":7,"display":"Three.js"},
    "linux administration":{"cat":"System","w":8,"display":"Linux Admin"},
    "web development":{"cat":"Frontend","w":7,"display":"Web Development"},
    "mobile development":{"cat":"Mobile","w":7,"display":"Mobile Development"},
    "software development":{"cat":"Programming","w":7,"display":"Software Development"},
    "full stack":{"cat":"Programming","w":8,"display":"Full Stack"},
    "full-stack":{"cat":"Programming","w":8,"display":"Full Stack"},
    "front end":{"cat":"Frontend","w":7,"display":"Frontend"},
    "front-end":{"cat":"Frontend","w":7,"display":"Frontend"},
    "back end":{"cat":"Backend","w":7,"display":"Backend"},
    "back-end":{"cat":"Backend","w":7,"display":"Backend"},
    "database management":{"cat":"Database","w":7,"display":"Database Management"},
    "cloud computing":{"cat":"Cloud","w":8,"display":"Cloud Computing"},
    "devops engineer":{"cat":"DevOps","w":8,"display":"DevOps"},
    "network security":{"cat":"Security","w":8,"display":"Network Security"},
    "information security":{"cat":"Security","w":8,"display":"Information Security"},
    "digital marketing":{"cat":"Marketing","w":6,"display":"Digital Marketing"},
    "search engine optimization":{"cat":"Marketing","w":6,"display":"SEO"},
    "content management":{"cat":"Management","w":5,"display":"Content Management"},
    "critical thinking":{"cat":"Soft Skills","w":8,"display":"Critical Thinking"},
    "time management":{"cat":"Soft Skills","w":7,"display":"Time Management"},
    "team leadership":{"cat":"Soft Skills","w":8,"display":"Team Leadership"},
    "verbal communication":{"cat":"Soft Skills","w":6,"display":"Communication"},
    "written communication":{"cat":"Soft Skills","w":6,"display":"Communication"},
    "tableau desktop":{"cat":"Data Viz","w":8,"display":"Tableau"},
    "ms excel":{"cat":"Productivity","w":6,"display":"Excel"},
    "ms word":{"cat":"Productivity","w":5,"display":"MS Word"},
    "google workspace":{"cat":"Productivity","w":5,"display":"Google Workspace"},
    "google analytics":{"cat":"Marketing","w":6,"display":"Google Analytics"},
    "amazon web services":{"cat":"Cloud","w":10,"display":"AWS"},
    "microsoft azure":{"cat":"Cloud","w":10,"display":"Azure"},
    "android studio":{"cat":"Mobile","w":8,"display":"Android Studio"},
    "xcode":{"cat":"Mobile","w":7,"display":"Xcode"},
    "visual studio":{"cat":"Productivity","w":6,"display":"Visual Studio"},
    "visual studio code":{"cat":"Productivity","w":6,"display":"VS Code"},
    "intellij":{"cat":"Productivity","w":6,"display":"IntelliJ"},
    "eclipse":{"cat":"Productivity","w":5,"display":"Eclipse"},
    "pycharm":{"cat":"Productivity","w":6,"display":"PyCharm"},
    "talend":{"cat":"Data Science","w":7,"display":"Talend"},
    "informatica":{"cat":"Data Science","w":7,"display":"Informatica"},
    "selenium webdriver":{"cat":"Testing","w":8,"display":"Selenium WebDriver"},
    "test automation":{"cat":"Testing","w":8,"display":"Test Automation"},
    "manual testing":{"cat":"Testing","w":6,"display":"Manual Testing"},
    "performance testing":{"cat":"Testing","w":7,"display":"Performance Testing"},
}

CERT_PATTERNS = [
    r'aws\s+certif\w+', r'azure\s+certif\w+', r'gcp\s+certif\w+',
    r'google\s+certif\w+', r'certified\s+\w+\s+\w+', r'\bpmp\b',
    r'\bcissp\b', r'\bceh\b', r'\bccna\b', r'\bccnp\b', r'\bitil\b',
    r'oracle\s+certif\w+', r'red\s+hat\s+certif\w+', r'\bcomptia\b',
    r'six\s+sigma', r'scrum\s+master', r'\bcsm\b', r'\bcka\b', r'\bckad\b',
    r'tensorflow\s+certif\w+', r'databricks\s+certif\w+',
    r'coursera\s+certif\w+', r'udemy\s+certif\w+', r'nptel\s+certif\w+',
    r'certif\w+\s+in\s+\w+', r'\w+\s+certif\w+\s+professional',
    r'\w+\s+certif\w+\s+associate', r'\w+\s+certif\w+\s+developer',
]

DEGREE_MAP = {
    r'\bb[\.\s]?tech\b':            'B.Tech',
    r'\bm[\.\s]?tech\b':            'M.Tech',
    r'\bb[\.\s]?e[\.\s]?\b':        'B.E.',
    r'\bb[\.\s]?sc\b':              'B.Sc',
    r'\bm[\.\s]?sc\b':              'M.Sc',
    r'\bb[\.\s]?c[\.\s]?a\b':       'BCA',
    r'\bm[\.\s]?c[\.\s]?a\b':       'MCA',
    r'\bb[\.\s]?b[\.\s]?a\b':       'BBA',
    r'\bm[\.\s]?b[\.\s]?a\b':       'MBA',
    r'\bph[\.\s]?d\b':              'PhD',
    r'\bb[\.\s]?com\b':             'B.Com',
    r'\bdiploma\b':                 'Diploma',
    r'\bbachelor\s+of\s+\w+':       'Bachelor\'s',
    r'\bmaster\s+of\s+\w+':         'Master\'s',
    r'\bsecondary\b|\b10th\b|\bssc\b': '10th / SSC',
    r'\bhigher\s+secondary\b|\b12th\b|\bhsc\b': '12th / HSC',
    r'\bbachelor':                  'Bachelor\'s Degree',
    r'\bmaster':                    'Master\'s Degree',
}

SENIORITY_MAP = {
    "intern":    ["intern","trainee","fresher","graduate trainee","student"],
    "junior":    ["junior","jr.","entry level","entry-level","associate developer"],
    "mid":       ["software engineer","developer","programmer","engineer","analyst"],
    "senior":    ["senior","sr.","lead","principal","staff engineer","tech lead"],
    "manager":   ["manager","team lead","head of","engineering manager","vp of"],
    "executive": ["director","cto","ceo","founder","co-founder","chief"],
}

SENIORITY_LABELS = {
    "intern":"Intern / Fresher","junior":"Junior (0-2 yrs)",
    "mid":"Mid-Level (2-5 yrs)","senior":"Senior (5-8 yrs)",
    "manager":"Team Lead / Manager","executive":"Director / Executive",
}

ACHIEVEMENT_PATTERNS = [
    r'\b\d+\s*%',
    r'[₹\$]\s*\d+',
    r'\b\d+\s*(lakh|crore|million|billion|thousand|k\b)',
    r'\bteam\s+of\s+\d+',
    r'\b\d+\s+(users|clients|customers|students|projects|applications)',
    r'\b(increased|improved|reduced|saved|achieved|delivered|built|launched|led|managed|optimized|developed|designed|implemented|created|deployed|automated)\b',
    r'\b(award|winner|rank|top|best|first|gold|silver|bronze|prize|honor|merit)\b',
    r'\bwon\b|\bselected\b|\brecognized\b|\bhonored\b',
]

# ══════════════════════════════════════════════════════════════
#  PDF TEXT EXTRACTION - 4 methods with fallbacks
# ══════════════════════════════════════════════════════════════
def extract_pdf_text(filepath):
    text = ""

    # Method 1: pdfplumber (best for modern PDFs)
    try:
        import pdfplumber
        debug("Trying pdfplumber...")
        with pdfplumber.open(filepath) as pdf:
            pages = []
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t)
            text = "\n".join(pages)
        if len(text.strip()) > 50:
            debug(f"pdfplumber OK: {len(text)} chars")
            return text
    except Exception as e:
        debug(f"pdfplumber failed: {e}")

    # Method 2: PyPDF2
    try:
        import PyPDF2
        debug("Trying PyPDF2...")
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            parts  = []
            for page in reader.pages:
                t = page.extract_text()
                if t: parts.append(t)
            text = "\n".join(parts)
        if len(text.strip()) > 50:
            debug(f"PyPDF2 OK: {len(text)} chars")
            return text
    except Exception as e:
        debug(f"PyPDF2 failed: {e}")

    # Method 3: pypdf (newer fork of PyPDF2)
    try:
        import pypdf
        debug("Trying pypdf...")
        with open(filepath, 'rb') as f:
            reader = pypdf.PdfReader(f)
            parts  = [p.extract_text() for p in reader.pages if p.extract_text()]
            text   = "\n".join(parts)
        if len(text.strip()) > 50:
            debug(f"pypdf OK: {len(text)} chars")
            return text
    except Exception as e:
        debug(f"pypdf failed: {e}")

    # Method 4: pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        debug("Trying pdfminer...")
        text = pdfminer_extract(filepath)
        if len(text.strip()) > 50:
            debug(f"pdfminer OK: {len(text)} chars")
            return text
    except Exception as e:
        debug(f"pdfminer failed: {e}")

    # Method 5: Raw binary extraction (last resort)
    try:
        debug("Trying raw binary extraction...")
        with open(filepath, 'rb') as f:
            raw = f.read()

        # Extract text between stream markers
        streams = re.findall(b'stream(.*?)endstream', raw, re.DOTALL)
        parts = []
        for s in streams:
            try:
                decoded = s.decode('latin-1', errors='ignore')
                # Get printable ASCII
                printable = re.sub(r'[^\x20-\x7E\n\t]', ' ', decoded)
                # Remove PDF operators
                cleaned = re.sub(r'\([^)]*\)', lambda m: m.group(0), printable)
                if len(cleaned.strip()) > 10:
                    parts.append(cleaned)
            except:
                pass

        text = "\n".join(parts)
        # Also try full binary as latin-1
        if len(text.strip()) < 50:
            text = re.sub(r'[^\x20-\x7E\n]', ' ', raw.decode('latin-1', errors='ignore'))

        debug(f"Raw binary: {len(text)} chars")
        return text

    except Exception as e:
        debug(f"Raw binary failed: {e}")

    return ""

# ══════════════════════════════════════════════════════════════
#  DOCX EXTRACTION
# ══════════════════════════════════════════════════════════════
def extract_docx_text(filepath):
    try:
        import docx
        doc  = docx.Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += " | ".join(c.text for c in row.cells) + "\n"
        return text
    except Exception as e:
        debug(f"python-docx failed: {e}")

    try:
        import zipfile
        with zipfile.ZipFile(filepath, 'r') as z:
            with z.open('word/document.xml') as f:
                xml = f.read().decode('utf-8', errors='ignore')
                text = re.sub(r'<[^>]+>', ' ', xml)
                return re.sub(r'\s+', ' ', text)
    except Exception as e:
        debug(f"DOCX zip fallback failed: {e}")
    return ""

def get_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    debug(f"Extracting: {filepath} (ext={ext})")

    if ext == '.pdf':
        return extract_pdf_text(filepath)
    elif ext == '.docx':
        return extract_docx_text(filepath)
    elif ext == '.doc':
        try:
            import subprocess
            r = subprocess.run(['antiword', filepath], capture_output=True, text=True, timeout=10)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout
        except Exception:
            pass
        try:
            with open(filepath, 'rb') as f:
                raw = f.read()
            return re.sub(r'[^\x20-\x7E\n\t]', ' ', raw.decode('latin-1', errors='ignore'))
        except:
            return ""
    return ""

# ══════════════════════════════════════════════════════════════
#  SKILL EXTRACTION - aggressive multi-pass matching
# ══════════════════════════════════════════════════════════════
def extract_skills(text):
    if not text:
        return [], {}

    # Normalize: lowercase, collapse spaces
    t = text.lower()
    t = re.sub(r'[•\-\*\|→▪◦►]', ' ', t)  # Remove bullets
    t = re.sub(r'[\(\)\[\]\{\}]', ' ', t)
    t = re.sub(r'[,;:/\\]', ' ', t)
    t = re.sub(r'\s+', ' ', t)

    found = {}  # display_name -> {cat, w}

    # Pass 1: Multi-word phrases (most important, check first)
    for phrase, meta in MULTI_WORD_SKILLS.items():
        if phrase in t:
            display = meta.get('display', phrase.title())
            found[display] = {"cat": meta['cat'], "w": meta['w']}
            debug(f"  multi-word match: '{phrase}' -> '{display}'")

    # Pass 2: Single-word skills with word boundary
    for skill, meta in SKILLS_DB.items():
        if ' ' not in skill:
            # Word boundary match (handles "react" not matching "reactjs" etc.)
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, t):
                # Choose display name
                if len(skill) <= 3:
                    display = skill.upper()
                elif skill in ('python','java','php','ruby','go','rust','swift','kotlin',
                               'scala','matlab','perl','dart','bash','linux','unix',
                               'docker','ansible','terraform','jenkins','jira','trello',
                               'figma','sketch','canva','selenium','cypress','playwright',
                               'postman','jupyter','tableau','pandas','numpy','spark',
                               'hadoop','kafka','redis','oracle','firebase','flutter',
                               'android','ios','ionic','xamarin','svelte','gatsby','nuxt',
                               'wordpress','shopify','magento','woocommerce',
                               'agile','scrum','kanban','confluence','bitbucket','svn',
                               'blockchain','ethereum','solidity'):
                    display = skill.capitalize()
                else:
                    display = skill.upper() if skill.isupper() else skill.title()

                if display not in found:  # Don't overwrite multi-word matches
                    found[display] = {"cat": meta['cat'], "w": meta['w']}
                    debug(f"  single-word match: '{skill}' -> '{display}'")

    # Pass 3: Special regex patterns
    # C++ / C#
    if re.search(r'\bc\+\+\b|c\s+plus\s+plus', text, re.I):
        found['C++'] = {"cat":"Programming","w":9}
    if re.search(r'\bc#\b', text, re.I):
        found['C#'] = {"cat":"Programming","w":8}
    # .NET
    if re.search(r'\.net\b', text, re.I):
        found['.NET'] = {"cat":"Backend","w":8}
    # R language (careful - only as standalone word)
    if re.search(r'\bR\b', text) and 'R' not in found:
        if re.search(r'\br\s+programming\b|\bprogramming\s+in\s+r\b|\br\s+language\b', t):
            found['R'] = {"cat":"Programming","w":7}
    # SQL variations
    if re.search(r'\bpl[\s\/]sql\b', t):
        found['PL/SQL'] = {"cat":"Database","w":7}
    if re.search(r'\bt[\s\-]sql\b', t):
        found['T-SQL'] = {"cat":"Database","w":7}
    # Check for "scikit" alone
    if re.search(r'\bscikit\b', t):
        found['Scikit-Learn'] = {"cat":"AI/ML","w":8}
    # Hugging Face
    if re.search(r'hugging\s*face|huggingface', t):
        found['Hugging Face'] = {"cat":"AI/ML","w":9}
    # VS Code
    if re.search(r'vs\s*code|visual\s+studio\s+code', t, re.I):
        found['VS Code'] = {"cat":"Productivity","w":6}
    # Power BI
    if re.search(r'power\s*bi', t):
        found['Power BI'] = {"cat":"Data Viz","w":8}
    # Spring Boot
    if re.search(r'spring\s*boot', t):
        found['Spring Boot'] = {"cat":"Backend","w":9}
    # React Native
    if re.search(r'react\s*native', t):
        found['React Native'] = {"cat":"Mobile","w":9}

    debug(f"Total skills found: {len(found)}")

    # Build categorized output
    cats = {}
    for name, meta in found.items():
        cat = meta['cat']
        if cat not in cats:
            cats[cat] = []
        cats[cat].append({"name": name, "w": meta['w']})

    for cat in cats:
        cats[cat].sort(key=lambda x: x['w'], reverse=True)

    # Flat list sorted by weight
    all_items = []
    for items in cats.values():
        all_items.extend(items)
    all_items.sort(key=lambda x: x['w'], reverse=True)
    flat = [i['name'] for i in all_items]

    by_cat = {cat: [s['name'] for s in skills] for cat, skills in cats.items()}

    return flat, by_cat

# ══════════════════════════════════════════════════════════════
#  CONTACT EXTRACTION
# ══════════════════════════════════════════════════════════════
def extract_contact(text):
    c = {"name":"","email":"","phone":"","linkedin":"","github":"","location":""}

    emails = re.findall(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', text)
    if emails: c['email'] = emails[0]

    phones = re.findall(r'(?:\+91[\-\s]?)?(?:\(?[0-9]{3,5}\)?[\-\s]?)?[0-9]{3,5}[\-\s]?[0-9]{4,5}', text)
    clean_phones = [re.sub(r'\D','', p) for p in phones if len(re.sub(r'\D','',p)) >= 10]
    if clean_phones: c['phone'] = clean_phones[0]

    lm = re.search(r'linkedin\.com/in/([A-Za-z0-9\-_]+)', text, re.I)
    if lm: c['linkedin'] = "linkedin.com/in/" + lm.group(1)

    gm = re.search(r'github\.com/([A-Za-z0-9\-_]+)', text, re.I)
    if gm: c['github'] = "github.com/" + gm.group(1)

    cities = ['mumbai','pune','delhi','bangalore','bengaluru','hyderabad','chennai',
              'kolkata','noida','gurugram','gurgaon','ahmedabad','jaipur','surat',
              'lucknow','nagpur','indore','bhopal','remote','new york','london',
              'singapore','dubai','san francisco','toronto']
    for city in cities:
        if re.search(r'\b' + city + r'\b', text, re.I):
            c['location'] = city.title(); break

    # Name: look for title-case 2-4 word sequence in first 10 lines
    for line in [l.strip() for l in text.split('\n') if l.strip()][:10]:
        line_clean = line.strip()
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$', line_clean):
            if not any(kw in line_clean.lower() for kw in
                       ['resume','curriculum','vitae','profile','objective','summary',
                        'university','college','institute','engineer','developer']):
                c['name'] = line_clean
                break

    return c

# ══════════════════════════════════════════════════════════════
#  EXPERIENCE
# ══════════════════════════════════════════════════════════════
def extract_experience(text):
    YEAR = r'\b(20\d{2}|19\d{2})\b'
    ROLE_KW = ['engineer','developer','analyst','manager','designer','consultant',
               'architect','intern','lead','director','scientist','specialist',
               'coordinator','officer','head','founder','associate','administrator',
               'researcher','programmer','trainee','executive','officer','hacker']

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    jobs  = []
    cur   = None

    for line in lines:
        ll = line.lower()
        has_role = any(kw in ll for kw in ROLE_KW)
        has_year = bool(re.search(YEAR, line))

        if has_role and len(line) < 120:
            if cur: jobs.append(cur)
            years = re.findall(YEAR, line)
            dur, yrs = "", 0
            if len(years) >= 2:
                try: yrs = int(years[1]) - int(years[0]); dur = f"{years[0]} – {years[1]}"
                except: pass
            elif years:
                dur = years[0] + " – Present"
                try: yrs = datetime.now().year - int(years[0])
                except: pass
            cur = {"title": line, "duration": dur, "years": yrs, "bullets": []}
        elif cur and len(line) > 15 and line.startswith(('•','–','-','*','→','►','◦','▪')) or (cur and len(line) > 20):
            cur["bullets"].append(line.lstrip('•–-*→►◦▪ ').strip())

    if cur: jobs.append(cur)

    # Total years from year range in full text
    all_years = [int(y) for y in re.findall(YEAR, text)]
    total = 0
    if all_years:
        mn, mx = min(all_years), min(max(all_years), datetime.now().year)
        total = mx - mn
        if total > 30: total = 0

    return jobs[:10], total

# ══════════════════════════════════════════════════════════════
#  EDUCATION
# ══════════════════════════════════════════════════════════════
def extract_education(text):
    tl = text.lower()
    degrees = []
    for pat, label in DEGREE_MAP.items():
        if re.search(pat, tl):
            # Get surrounding context
            m = re.search(pat, tl)
            if m:
                s = max(0, m.start()-20)
                e = min(len(text), m.end()+80)
                ctx = re.sub(r'\s+',' ', text[s:e]).strip()
                degrees.append({"degree": label, "context": ctx})

    insts = re.findall(
        r'([A-Z][A-Za-z\s,\.]+(?:university|college|institute|school|iit|nit|bits|vit|srm|manipal|amity)[A-Za-z\s,\.]*)',
        text, re.I)

    gpa_m = re.search(r'(cgpa|gpa|percentage|marks|score)[:\s]+(\d+\.?\d*)', text, re.I)
    gpa   = gpa_m.group(0).strip() if gpa_m else ""

    years = sorted(set(re.findall(r'\b(20\d{2}|19\d{2})\b', text)))

    return {"degrees": degrees[:4], "institutions": list(set(insts[:3])),
            "gpa": gpa, "years": years}

# ══════════════════════════════════════════════════════════════
#  PROJECTS
# ══════════════════════════════════════════════════════════════
def extract_projects(text):
    # Find project section
    proj_section = ""
    m = re.search(
        r'(project|portfolio|work\s+done)[^\n]*\n(.*?)(?=\n(?:education|experience|skill|certif|achieve|language|reference|contact|objective|summary)|$)',
        text, re.I | re.DOTALL)
    if m:
        proj_section = m.group(2)
    else:
        proj_section = text  # scan full text

    lines    = [l.strip() for l in proj_section.split('\n') if l.strip()]
    projects = []
    cur      = None

    for line in lines:
        if len(line) < 8: continue
        # A project title is usually short, non-bullet, no period at end
        is_title = (len(line) < 80 and not line.endswith('.')
                    and not line.startswith(('•','–','-','*'))
                    and not re.match(r'^\d{4}', line))
        if is_title and cur is None:
            cur = {"title": line, "desc": [], "techs": []}
        elif cur:
            cur["desc"].append(line.lstrip('•–-* '))
            # Detect tech in description
            s,_ = extract_skills(line)
            cur["techs"].extend(s[:4])
        if len(projects) >= 8 and cur:
            projects.append(cur); cur = None; break

    if cur: projects.append(cur)

    for p in projects:
        p["techs"] = list(dict.fromkeys(p["techs"]))[:8]

    return projects[:8]

# ══════════════════════════════════════════════════════════════
#  CERTIFICATIONS
# ══════════════════════════════════════════════════════════════
def extract_certifications(text):
    tl   = text.lower()
    certs = []
    for pat in CERT_PATTERNS:
        for m in re.finditer(pat, tl):
            s   = max(0, m.start()-10)
            e   = min(len(text), m.end()+80)
            ctx = re.sub(r'\s+',' ', text[s:e]).strip()
            ctx = ctx.lstrip('•–-* ')
            if ctx and len(ctx) > 8 and ctx not in certs:
                certs.append(ctx)
    return list(dict.fromkeys(certs))[:10]

# ══════════════════════════════════════════════════════════════
#  ACHIEVEMENTS
# ══════════════════════════════════════════════════════════════
def extract_achievements(text):
    results = []
    for line in text.split('\n'):
        line = line.strip().lstrip('•–-*→►◦▪ ')
        if len(line) < 20: continue
        score = sum(1 for p in ACHIEVEMENT_PATTERNS if re.search(p, line, re.I))
        if score >= 1:
            results.append({"t": line, "s": score})
    results.sort(key=lambda x: x['s'], reverse=True)
    return [r['t'] for r in results[:8]]

# ══════════════════════════════════════════════════════════════
#  SENIORITY
# ══════════════════════════════════════════════════════════════
def get_seniority(text, total_exp):
    tl     = text.lower()
    scores = {k: 0 for k in SENIORITY_MAP}

    for level, kws in SENIORITY_MAP.items():
        for kw in kws:
            if kw in tl: scores[level] += 1

    if total_exp == 0:    scores['intern']  += 5
    elif total_exp <= 1:  scores['junior']  += 5
    elif total_exp <= 3:  scores['mid']     += 5
    elif total_exp <= 6:  scores['senior']  += 4
    elif total_exp <= 10: scores['manager'] += 3
    else:                 scores['executive'] += 3

    best = max(scores, key=scores.get)
    return best, SENIORITY_LABELS.get(best, "Mid-Level")

# ══════════════════════════════════════════════════════════════
#  LANGUAGES
# ══════════════════════════════════════════════════════════════
def get_human_languages(text):
    langs = ["english","hindi","marathi","gujarati","bengali","tamil","telugu",
             "kannada","malayalam","urdu","punjabi","french","german","spanish",
             "arabic","japanese","chinese","mandarin","italian","portuguese","korean"]
    tl = text.lower()
    return [l.title() for l in langs if re.search(r'\b'+l+r'\b', tl)]

# ══════════════════════════════════════════════════════════════
#  SCORE
# ══════════════════════════════════════════════════════════════
def calc_score(flat, edu, exp_jobs, certs, achievements, projects, total_exp, github, linkedin):
    pts = {}
    pts['skills']         = min(len(flat) * 1.8, 30)
    pts['experience']     = min(total_exp * 3, 25)
    pts['education']      = (5 if edu['gpa'] else 0) + (5 if edu['degrees'] else 0) + (5 if edu['institutions'] else 0)
    pts['projects']       = min(len(projects) * 2.5, 10)
    pts['certifications'] = min(len(certs) * 2.5, 10)
    pts['achievements']   = min(len(achievements), 5)
    pts['online_presence']= (2.5 if github else 0) + (2.5 if linkedin else 0)
    total = round(min(sum(pts.values()), 100))
    return total, {k: round(v) for k,v in pts.items()}

# ══════════════════════════════════════════════════════════════
#  AI SUMMARY
# ══════════════════════════════════════════════════════════════
def make_summary(contact, flat, by_cat, total_exp, seniority, edu, certs, score, achievements):
    name    = contact.get('name') or 'The Candidate'
    top5    = ", ".join(flat[:5]) if flat else "various technologies"
    exp_str = ("a fresher / recent graduate" if total_exp==0 else
               f"~{total_exp} year{'s' if total_exp!=1 else ''} of experience")

    domain_map = {"AI/ML":"AI & Machine Learning","Frontend":"Frontend Development",
                  "Backend":"Backend Development","Data Science":"Data Science",
                  "DevOps":"DevOps & Cloud","Mobile":"Mobile Development",
                  "Security":"Cybersecurity","Blockchain":"Blockchain","Programming":"Software Engineering"}
    primary = "Software Development"
    if by_cat:
        top_cat = max(by_cat, key=lambda c: len(by_cat[c]))
        primary = domain_map.get(top_cat, top_cat)

    edu_note  = f" Holds a {edu['degrees'][0]['degree']}." if edu.get('degrees') else ""
    cert_note = f" Has {len(certs)} certification(s)." if certs else ""
    ach_note  = f" Notable: {achievements[0][:70]}." if achievements else ""
    loc_note  = f" Based in {contact['location']}." if contact.get('location') else ""

    return (f"{name} is a {seniority} specializing in {primary} with {exp_str}. "
            f"Key skills: {top5}.{edu_note}{cert_note}{ach_note}{loc_note} "
            f"Profile score: {score}/100.")

# ══════════════════════════════════════════════════════════════
#  RED FLAGS & STRENGTHS
# ══════════════════════════════════════════════════════════════
def get_flags(text, total_exp, flat, contact):
    flags = []
    if len(flat) < 4:
        flags.append("Very few technical skills detected — add a dedicated Skills section to your resume")
    if not contact.get('email'):
        flags.append("No email address found in resume")
    if not contact.get('phone'):
        flags.append("No phone number found in resume")
    if len(text.split()) < 100:
        flags.append("Resume appears very short — consider adding more detail")
    if total_exp == 0 and not re.search(r'intern|project|training|course', text, re.I):
        flags.append("No experience or projects mentioned — add internships or personal projects")
    years = sorted(set(int(y) for y in re.findall(r'\b(20\d{2})\b', text)))
    for i in range(len(years)-1):
        if years[i+1] - years[i] > 2:
            flags.append(f"Possible employment/study gap detected around {years[i]}–{years[i+1]}")
            break
    return flags

def get_strengths(by_cat, certs, achievements, contact, total_exp, projects):
    s = []
    hd = {"AI/ML","Cloud","DevOps","Blockchain","Security"}
    for cat in by_cat:
        if cat in hd: s.append(f"High-demand {cat} skills present")
    if len(by_cat.get('Programming',[])) >= 3:
        s.append("Polyglot programmer — multiple programming languages")
    if certs:
        s.append(f"{len(certs)} professional certifications detected")
    if achievements:
        s.append("Quantifiable achievements with measurable impact")
    if contact.get('github'):
        s.append("Active GitHub profile — hands-on coding evidence")
    if contact.get('linkedin'):
        s.append("LinkedIn professional presence")
    if len(projects) >= 3:
        s.append(f"Strong project portfolio — {len(projects)} projects")
    if total_exp >= 5:
        s.append(f"{total_exp}+ years of industry experience")
    return s

# ══════════════════════════════════════════════════════════════
#  JOB MATCH
# ══════════════════════════════════════════════════════════════
def job_match(flat, job_csv):
    if not job_csv: return None
    job_skills   = [s.strip().lower() for s in job_csv.split(',') if s.strip()]
    resume_lower = [s.lower() for s in flat]
    matched, missing = [], []
    for js in job_skills:
        found = any(js in rs or rs in js for rs in resume_lower)
        (matched if found else missing).append(js)
    pct = round(len(matched)/len(job_skills)*100, 1) if job_skills else 0
    return {"percentage": pct, "matched_skills": matched, "missing_skills": missing,
            "total_required": len(job_skills), "total_matched": len(matched),
            "recommendation": ("Strong match — shortlist" if pct>=70 else
                               "Moderate match — review" if pct>=40 else "Weak match")}

# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════
def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python resume_parser.py <filepath>", "skills":"[]"}))
        sys.exit(1)

    filepath = sys.argv[1]
    job_skills_csv = sys.argv[2] if len(sys.argv) > 2 else ""

    if not os.path.exists(filepath):
        print(json.dumps({"error": "File not found", "skills": "[]"}))
        sys.exit(1)

    debug(f"=== RecruitAI Parser v3 ===")
    debug(f"File: {filepath}")

    text = get_text(filepath)
    debug(f"Extracted text length: {len(text)} chars")
    debug(f"First 200 chars: {repr(text[:200])}")

    if not text or len(text.strip()) < 30:
        print(json.dumps({
            "success": False,
            "error": "Could not extract text. This PDF may be scanned/image-based.",
            "skills": "[]",
            "skills_flat": [],
            "total_skills": 0
        }))
        sys.exit(0)

    # 🔍 Run extractors
    flat, by_cat = extract_skills(text)
    contact = extract_contact(text)
    exp_jobs, total_exp = extract_experience(text)
    education = extract_education(text)
    projects = extract_projects(text)
    certs = extract_certifications(text)
    achievements = extract_achievements(text)
    languages = get_human_languages(text)
    seniority_key, slabel = get_seniority(text, total_exp)
    score, breakdown = calc_score(
        flat, education, exp_jobs, certs, achievements,
        projects, total_exp,
        contact.get('github'), contact.get('linkedin')
    )

    result = {
    "success": True,

    "skills": flat,
    "skills_flat": flat,
    "total_skills": len(flat),
    "skills_by_category": by_cat,

    "contact": contact,
    "experience": exp_jobs,
    "education": education,
    "projects": projects,
    "certifications": certs,
    "achievements": achievements,
    "languages": languages,

    "score": score,

    "word_count": len(text.split()),
    "text_preview": text[:2000]
}

    # ✅ FINAL OUTPUT (VERY IMPORTANT)
    print(json.dumps(result, ensure_ascii=False))


# ✅ MUST BE AT VERY BOTTOM (NO INDENT)
if __name__ == "__main__":
    main()